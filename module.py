from docx import Document



def word(mavzu, text, id):
    document = Document()
    p = document.add_paragraph(mavzu)
    p.alignment = 1
    p1 = document.add_paragraph(text)
    document.save(f'demo{id}.docx')
    return f'demo{id}.docx'